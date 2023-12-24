import os
import re
import docx
import openpyxl
import pdfplumber


def findPhone(value):

    if type(value) != str:
        value = str(value)
    
    return re.findall("(1[3-9]\d{9})", value)


def getPhoneFromPdf(pdfPath):
    result = []
    pdf = pdfplumber.open(pdfPath)

    for page in pdf.pages:
        text = page.extract_text()
        result.extend(findPhone(text))
        tables = page.extract_tables()

        for table in tables:
            for row in table:
                for cell in row:
                    result.extend(findPhone(cell))

    return result


def getPhoneFromWord(wordPath):
    result = []
    doc = docx.Document(wordPath)

    for p in doc.paragraphs:
        result.extend(findPhone(p.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                result.extend(findPhone(cell.text))
    
    return result


def getPhoneFromExcel(excelPath):
    result = []
    wb = openpyxl.load_workbook(excelPath)

    for sheet in wb.worksheets:
        for row in sheet.rows:
            for cell in row:
                result.extend(findPhone(cell.value)) 

    return result


def getPhoneFromFiles(filePath):

    allItems = os.listdir(filePath)
    

    result = []
    
    for item in allItems:
        itemPath = os.path.join(filePath, item)
        
        extension = os.path.splitext(item)[1].lower()
        
        if os.path.isdir(itemPath):
            result.extend(getPhoneFromFiles(itemPath))

        elif extension in [".pdf"]:
            pdfResult = getPhoneFromPdf(itemPath)
            result.extend(pdfResult)

        elif extension in [".docx"]:
            wordResult = getPhoneFromWord(itemPath)
            result.extend(wordResult)

        elif extension in [".xlsx"]:
            excelResult = getPhoneFromExcel(itemPath)
            result.extend(excelResult)
    
    return result


phoneList = getPhoneFromFiles("/Users/yequ/misc")

phoneList = list(set(phoneList))

print(phoneList)